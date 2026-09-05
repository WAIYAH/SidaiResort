#!/usr/bin/env python
"""Generate the one-page shareable menu card (.docx) from data/menu.json.

    python tools/build-menu-card.py

Same source of truth as the website, so the card cannot drift from the page.
Prints on a single A4 sheet: logo, two columns of sections, contact footer.

Design follows the printed menu: navy #0F214A and amber #F2B824, with the
resort green #1A4D2E alternating on the section bars. The page itself is left
white rather than navy - Word does not print background colours by default, so
a navy page would come out of most printers blank behind the type.
"""
import io
import json
import os
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA = os.path.join(ROOT, "data", "menu.json")
LOGO = os.path.join(ROOT, "assets", "images", "branding", "sidai-logo.png")
OUT = os.path.join(ROOT, "assets", "downloads", "Sidai-Resort-Menu-Card.docx")

NAVY = RGBColor(0x0F, 0x21, 0x4A)
GOLD = RGBColor(0xF2, 0xB8, 0x24)
GREEN = RGBColor(0x1A, 0x4D, 0x2E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x5A, 0x63, 0x72)

PHONE = "0703 761 951 / 0705 800 802"
EMAIL = "sidairesort21@gmail.com"

# Split so the two columns end at roughly the same depth.
LEFT = ["breakfast", "local", "accompaniments", "desserts"]
RIGHT = ["nyama-choma", "snacks", "drinks", "hot", "beers"]

COL_W = Inches(3.5)


def shade(el, hex_fill):
    """Fill a paragraph or cell with a solid colour."""
    pr = el.get_or_add_pPr() if el.tag.endswith("}p") else el.get_or_add_tcPr()  # noqa: E501
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:fill"), hex_fill)
    pr.append(s)


def fixed_layout(table, width_twips):
    """Word ignores cell.width unless the table layout is fixed and the grid is
    declared, which left one column's shaded bars short of the full width."""
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    grid = table._tbl.find(qn("w:tblGrid"))
    for gc in list(grid):
        grid.remove(gc)
    for _ in range(2):
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(width_twips))
        grid.append(gc)


def no_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        borders.append(e)
    tbl_pr.append(borders)


def rule(par, hex_color, size=8):
    """Draw a horizontal line under a paragraph."""
    pr = par._element.get_or_add_pPr()
    bd = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size))
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), hex_color)
    bd.append(b)
    pr.append(bd)


def money(price):
    fmt = lambda n: f"{float(n):,.2f}"
    if isinstance(price, (int, float)):
        return fmt(price)
    return " / ".join(fmt(p.strip()) for p in str(price).split("/"))


def run(par, text, *, size, color, bold=False, italic=False, caps=False, space=0):
    r = par.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    if caps or space:
        rpr = r._element.get_or_add_rPr()
        if caps:
            c = OxmlElement("w:caps")
            rpr.append(c)
        if space:
            sp = OxmlElement("w:spacing")
            sp.set(qn("w:val"), str(int(space * 20)))
            rpr.append(sp)
    return r


def tight(par, before=0, after=0, line=None):
    pf = par.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)
    return par


def build():
    data = json.load(io.open(DATA, encoding="utf-8"))
    sections = {s["id"]: s for s in data["sections"]}
    doc = Document()

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)      # A4
    sec.left_margin = sec.right_margin = Inches(0.45)
    sec.top_margin = Inches(0.35)
    sec.bottom_margin = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(8)
    normal.paragraph_format.space_after = Pt(0)

    # ---------- masthead ----------
    p = tight(doc.add_paragraph(), after=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(LOGO, height=Inches(0.5))

    p = tight(doc.add_paragraph(), after=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, "SIDAI RESORT & HOTEL", size=19, color=NAVY, bold=True, space=1.4)

    p = tight(doc.add_paragraph(), after=4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, "FOOD & BEVERAGE MENU", size=9, color=GREEN, bold=True, space=2.6)

    p = tight(doc.add_paragraph(), after=6)
    rule(p, "F2B824", size=12)

    # ---------- two columns ----------
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(table)
    table.autofit = False
    fixed_layout(table, int(COL_W.inches * 1440))
    cells = table.rows[0].cells
    for cell, ids in zip(cells, (LEFT, RIGHT)):
        cell.width = COL_W
        cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
        for n, sid in enumerate(ids):
            s = sections[sid]
            navy_bar = (list(sections).index(sid) % 2 == 0)

            # section bar
            p = cell.add_paragraph()
            tight(p, before=(7 if n else 0), after=2, line=11)
            p.paragraph_format.left_indent = Pt(0)
            shade(p._element, "0F214A" if navy_bar else "1A4D2E")
            run(p, "  " + s["title"], size=8.5, color=(GOLD if navy_bar else WHITE),
                bold=True, caps=True, space=1.0)

            # subtitle
            p = tight(cell.add_paragraph(), after=3, line=9)
            run(p, "  " + s["subtitle"], size=7, color=GREY, italic=True)

            # items, price right-aligned on a dotted leader
            for item in s["items"]:
                p = cell.add_paragraph()
                tight(p, after=0.5, line=10.5)
                stops = p.paragraph_format.tab_stops
                stops.add_tab_stop(COL_W - Inches(0.12), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
                run(p, item["name"], size=8, color=NAVY, bold=True)
                run(p, "\t" + money(item["price"]), size=8, color=GREEN, bold=True)

    # ---------- footer ----------
    p = tight(doc.add_paragraph(), before=8, after=2)
    rule(p, "F2B824", size=12)

    p = tight(doc.add_paragraph(), after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, "Sidai Resort & Hotel", size=8, color=NAVY, bold=True)
    run(p, "  |  Orders & Pre-Orders: ", size=8, color=GREY)
    run(p, PHONE, size=8, color=NAVY, bold=True)
    run(p, "  |  ", size=8, color=GREY)
    run(p, EMAIL, size=8, color=NAVY, bold=True)

    p = tight(doc.add_paragraph(), before=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, "FRCX+PP4 Naroosura, Narok County, Kenya   ·   sidairesort.com   ·   All prices in Kenya Shillings (KSh)",
        size=7, color=GREY)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    n = sum(len(s["items"]) for s in data["sections"])
    print(f"{OUT}\n  {len(data['sections'])} sections, {n} items")
    return OUT


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    build()
